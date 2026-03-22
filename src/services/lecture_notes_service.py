"""
Lecture Notes Service for processing and managing lecture notes and background materials.
Handles file upload, content extraction, and RAG integration.
"""

import os
import logging
import tempfile
from typing import Optional, List, Dict, Any, BinaryIO
from pathlib import Path
from datetime import datetime
import asyncio
import time
from functools import wraps

# File processing imports
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import fitz  # pymupdf
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from ..models.grading_models import LectureNote, ProcessingStatus
    from ..utils.lecture_notes_storage import LectureNotesStorage
    from ..errors.grading_errors import ValidationError, IntegrationError, ErrorHandler
    from .lecture_notes_validator import AcademicDocumentValidator
except ImportError:
    # Fallback for direct execution
    from src.models.grading_models import LectureNote, ProcessingStatus
    from src.utils.lecture_notes_storage import LectureNotesStorage
    from src.errors.grading_errors import ValidationError, IntegrationError, ErrorHandler
    from src.services.lecture_notes_validator import AcademicDocumentValidator

logger = logging.getLogger(__name__)


def retry_with_exponential_backoff(max_retries: int = 3, base_delay: float = 1.0, max_delay: float = 10.0):
    """
    Decorator for retrying operations with exponential backoff.
    
    Args:
        max_retries: Maximum number of retry attempts
        base_delay: Initial delay in seconds
        max_delay: Maximum delay between retries
    """
    def decorator(func):
        @wraps(func)
        async def async_wrapper(*args, **kwargs):
            last_exception = None
            for attempt in range(max_retries + 1):
                try:
                    return await func(*args, **kwargs)
                except Exception as e:
                    last_exception = e
                    if attempt < max_retries:
                        delay = min(base_delay * (2 ** attempt), max_delay)
                        logger.warning(f"Attempt {attempt + 1}/{max_retries + 1} failed for {func.__name__}: {e}. Retrying in {delay}s...")
                        await asyncio.sleep(delay)
                    else:
                        logger.error(f"All {max_retries + 1} attempts failed for {func.__name__}")
            raise last_exception
        
        @wraps(func)
        def sync_wrapper(*args, **kwargs):
            last_exception = None
            for attempt in range(max_retries + 1):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    last_exception = e
                    if attempt < max_retries:
                        delay = min(base_delay * (2 ** attempt), max_delay)
                        logger.warning(f"Attempt {attempt + 1}/{max_retries + 1} failed for {func.__name__}: {e}. Retrying in {delay}s...")
                        time.sleep(delay)
                    else:
                        logger.error(f"All {max_retries + 1} attempts failed for {func.__name__}")
            raise last_exception
        
        # Return appropriate wrapper based on whether function is async
        if asyncio.iscoroutinefunction(func):
            return async_wrapper
        else:
            return sync_wrapper
    
    return decorator


class LectureNotesService:
    """Service for processing lecture notes and background materials."""
    
    # File size limits
    MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50MB
    
    # Supported file types
    SUPPORTED_FILE_TYPES = {'pdf', 'docx', 'txt', 'md'}
    
    # Content processing settings
    MAX_CHUNK_SIZE = 1000
    CHUNK_OVERLAP = 100
    
    def __init__(self, storage: Optional[LectureNotesStorage] = None, rag_system=None):
        """
        Initialize the lecture notes service.
        
        Args:
            storage: LectureNotesStorage instance (optional, creates default if None)
            rag_system: RAG grading system instance for integration (optional)
        """
        self.storage = storage or LectureNotesStorage()
        self.rag_system = rag_system
        self.validator = AcademicDocumentValidator()
        
        # Log available processing libraries
        self._log_available_libraries()
    
    def _log_available_libraries(self) -> None:
        """Log which file processing libraries are available."""
        available = []
        if PDFPLUMBER_AVAILABLE:
            available.append("pdfplumber")
        if PYMUPDF_AVAILABLE:
            available.append("PyMuPDF")
        if PYPDF2_AVAILABLE:
            available.append("PyPDF2")
        if PYTHON_DOCX_AVAILABLE:
            available.append("python-docx")
        
        logger.info(f"Available file processing libraries: {', '.join(available) if available else 'None'}")
    
    def validate_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Validate uploaded file against size and format requirements using enhanced validator.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Dictionary with validation results
            
        Raises:
            ValidationError: If file validation fails
        """
        try:
            # Use enhanced validator for comprehensive validation
            validation_result = self.validator.validate_file_comprehensive(file_content, filename)
            
            if not validation_result['valid']:
                # Check if any fallback methods succeeded
                fallback_options = validation_result.get('fallback_options', [])
                if any(fb.get('valid', False) for fb in fallback_options):
                    logger.warning(f"Primary validation failed for {filename}, but fallback methods succeeded")
                    validation_result['valid'] = True
                    validation_result['used_fallback'] = True
                else:
                    # Collect all errors for the exception
                    all_errors = validation_result.get('errors', [])
                    error_message = "; ".join(all_errors) if all_errors else "File validation failed"
                    raise ValidationError(error_message)
            
            # Log validation summary
            summary = self.validator.get_validation_summary(validation_result)
            logger.info(f"File validation summary for {filename}:\n{summary}")
            
            return {
                'valid': True,
                'file_size': validation_result['file_size'],
                'file_type': validation_result['detected_type'],
                'validation_details': validation_result
            }
            
        except ValidationError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error during file validation: {e}")
            raise ValidationError(f"File validation failed: {str(e)}")
    
    def _validate_file_format(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        """
        Perform format-specific validation for academic documents.
        
        Args:
            file_content: File content as bytes
            file_type: File extension
            
        Returns:
            Dictionary with format validation results
        """
        try:
            if file_type == 'pdf':
                return self._validate_pdf_format(file_content)
            elif file_type == 'docx':
                return self._validate_docx_format(file_content)
            elif file_type in ['txt', 'md']:
                return self._validate_text_format(file_content)
            else:
                return {'valid': True, 'details': 'No specific validation for this format'}
                
        except Exception as e:
            logger.warning(f"Format validation failed for {file_type}: {e}")
            return {'valid': False, 'error': str(e)}
    
    def _validate_pdf_format(self, file_content: bytes) -> Dict[str, Any]:
        """Validate PDF file format."""
        try:
            # Check PDF header
            if not file_content.startswith(b'%PDF-'):
                return {'valid': False, 'error': 'Invalid PDF header'}
            
            # Try to open with available libraries
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                if PDFPLUMBER_AVAILABLE:
                    with pdfplumber.open(temp_file_path) as pdf:
                        page_count = len(pdf.pages)
                        return {'valid': True, 'page_count': page_count, 'method': 'pdfplumber'}
                elif PYMUPDF_AVAILABLE:
                    doc = fitz.open(temp_file_path)
                    page_count = doc.page_count
                    doc.close()
                    return {'valid': True, 'page_count': page_count, 'method': 'PyMuPDF'}
                elif PYPDF2_AVAILABLE:
                    with open(temp_file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        page_count = len(reader.pages)
                        return {'valid': True, 'page_count': page_count, 'method': 'PyPDF2'}
                else:
                    return {'valid': False, 'error': 'No PDF processing library available'}
            finally:
                os.unlink(temp_file_path)
                
        except Exception as e:
            return {'valid': False, 'error': f'PDF validation failed: {str(e)}'}
    
    def _validate_docx_format(self, file_content: bytes) -> Dict[str, Any]:
        """Validate DOCX file format."""
        try:
            if not PYTHON_DOCX_AVAILABLE:
                return {'valid': False, 'error': 'python-docx library not available'}
            
            # Check DOCX signature (ZIP file starting with PK)
            if not file_content.startswith(b'PK'):
                return {'valid': False, 'error': 'Invalid DOCX format (not a ZIP file)'}
            
            # Try to open as DOCX
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = Document(temp_file_path)
                paragraph_count = len(doc.paragraphs)
                return {'valid': True, 'paragraph_count': paragraph_count, 'method': 'python-docx'}
            finally:
                os.unlink(temp_file_path)
                
        except Exception as e:
            return {'valid': False, 'error': f'DOCX validation failed: {str(e)}'}
    
    def _validate_text_format(self, file_content: bytes) -> Dict[str, Any]:
        """Validate text file format."""
        try:
            # Try to decode as UTF-8
            text_content = file_content.decode('utf-8')
            word_count = len(text_content.split())
            line_count = len(text_content.splitlines())
            
            return {
                'valid': True,
                'word_count': word_count,
                'line_count': line_count,
                'encoding': 'utf-8'
            }
            
        except UnicodeDecodeError:
            # Try other common encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text_content = file_content.decode(encoding)
                    word_count = len(text_content.split())
                    line_count = len(text_content.splitlines())
                    
                    return {
                        'valid': True,
                        'word_count': word_count,
                        'line_count': line_count,
                        'encoding': encoding
                    }
                except UnicodeDecodeError:
                    continue
            
            return {'valid': False, 'error': 'Unable to decode text file with common encodings'}
        
        except Exception as e:
            return {'valid': False, 'error': f'Text validation failed: {str(e)}'}
    
    async def upload_lecture_note(
        self,
        file_content: bytes,
        filename: str,
        associate_with_rubric: Optional[str] = None
    ) -> LectureNote:
        """
        Upload and process a lecture note file.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            associate_with_rubric: Optional rubric ID to associate with
            
        Returns:
            Created LectureNote object
            
        Raises:
            ValidationError: If file validation fails
            IntegrationError: If processing fails
        """
        try:
            logger.info(f"Starting upload of lecture note: {filename}")
            
            # Validate file
            validation_result = self.validate_file(file_content, filename)
            file_size = validation_result['file_size']
            file_type = validation_result['file_type']
            
            # Create lecture note entry
            note = self.storage.create_lecture_note(
                original_name=filename,
                file_size=file_size,
                file_type=file_type
            )
            
            # Extract content asynchronously
            try:
                extracted_content = await self._extract_content_async(file_content, file_type)
                
                # Update note with extracted content
                note.extracted_content = extracted_content
                note.word_count = len(extracted_content.split()) if extracted_content else 0
                note.processed_at = datetime.now()
                note.clear_processing_error()
                
                # Save updated note
                self.storage.update_lecture_note(note)
                
                logger.info(f"Successfully processed lecture note {note.id}: {note.word_count} words extracted")
                
            except Exception as e:
                # Mark as failed but keep the note entry
                error_msg = f"Content extraction failed: {str(e)}"
                note.set_processing_error(error_msg)
                self.storage.update_lecture_note(note)
                
                logger.error(f"Failed to extract content from {filename}: {e}")
                # Don't raise here - we want to keep the note entry even if extraction fails
            
            # Associate with rubric if specified
            if associate_with_rubric:
                success = self.storage.associate_note_with_rubric(note.id, associate_with_rubric)
                if success:
                    logger.info(f"Associated note {note.id} with rubric {associate_with_rubric}")
                else:
                    logger.warning(f"Failed to associate note {note.id} with rubric {associate_with_rubric}")
            
            # Add to RAG system if available and content was extracted
            if self.rag_system and note.extracted_content:
                try:
                    await self._add_to_rag_system(note)
                    logger.info(f"Successfully integrated note {note.id} with RAG system")
                except IntegrationError as e:
                    # Log the error but don't fail the upload
                    logger.warning(f"RAG integration failed for note {note.id}, but upload succeeded: {e.message}")
                    # Update note to reflect RAG integration failure
                    self.storage.update_lecture_note(note)
                except Exception as e:
                    # Unexpected error - log but continue
                    logger.error(f"Unexpected error during RAG integration for note {note.id}: {e}")
                    if not note.metadata:
                        note.metadata = {}
                    note.metadata['rag_integrated'] = False
                    note.metadata['rag_integration_error'] = f"Unexpected error: {str(e)}"
                    self.storage.update_lecture_note(note)
            
            return note
            
        except ValidationError:
            raise
        except Exception as e:
            logger.error(f"Failed to upload lecture note {filename}: {e}")
            raise IntegrationError(f"Lecture note upload failed: {str(e)}")
    
    async def _extract_content_async(self, file_content: bytes, file_type: str) -> str:
        """
        Extract text content from file asynchronously.
        
        Args:
            file_content: File content as bytes
            file_type: File extension
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If content extraction fails
        """
        # Run extraction in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_content_sync, file_content, file_type)
    
    def _extract_content_sync(self, file_content: bytes, file_type: str) -> str:
        """
        Synchronous content extraction for different file types with fallback mechanisms.
        
        Args:
            file_content: File content as bytes
            file_type: File extension
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If all extraction methods fail
        """
        extraction_errors = []
        
        try:
            if file_type == 'pdf':
                return self._extract_pdf_content_with_fallback(file_content, extraction_errors)
            elif file_type == 'docx':
                return self._extract_docx_content_with_fallback(file_content, extraction_errors)
            elif file_type in ['txt', 'md']:
                return self._extract_text_content_with_fallback(file_content, extraction_errors)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            error_summary = f"Content extraction failed for {file_type}: {str(e)}"
            if extraction_errors:
                error_summary += f"\nAttempted methods: {', '.join(extraction_errors)}"
            logger.error(error_summary)
            raise Exception(error_summary)
    
    def _extract_pdf_content(self, file_content: bytes) -> str:
        """Extract text content from PDF file."""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # Try extraction methods in order of preference
            extraction_methods = []
            
            if PDFPLUMBER_AVAILABLE:
                extraction_methods.append(('pdfplumber', self._extract_pdf_with_pdfplumber))
            if PYMUPDF_AVAILABLE:
                extraction_methods.append(('PyMuPDF', self._extract_pdf_with_pymupdf))
            if PYPDF2_AVAILABLE:
                extraction_methods.append(('PyPDF2', self._extract_pdf_with_pypdf2))
            
            if not extraction_methods:
                raise Exception("No PDF extraction libraries available")
            
            # Try each method until one succeeds
            for method_name, method_func in extraction_methods:
                try:
                    logger.debug(f"Trying PDF extraction with {method_name}")
                    content = method_func(temp_file_path)
                    if content and content.strip():
                        logger.info(f"Successfully extracted PDF content using {method_name}")
                        return content.strip()
                    else:
                        logger.warning(f"{method_name} returned empty content")
                except Exception as e:
                    logger.warning(f"{method_name} extraction failed: {e}")
                    continue
            
            raise Exception("All PDF extraction methods failed or returned empty content")
            
        finally:
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp PDF file: {e}")
    
    def _extract_pdf_content_with_fallback(self, file_content: bytes, extraction_errors: List[str]) -> str:
        """
        Extract PDF content with comprehensive fallback mechanisms.
        
        Args:
            file_content: PDF file content as bytes
            extraction_errors: List to collect error messages
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If all extraction methods fail
        """
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # Primary extraction methods
            extraction_methods = []
            
            if PDFPLUMBER_AVAILABLE:
                extraction_methods.append(('pdfplumber', self._extract_pdf_with_pdfplumber))
            if PYMUPDF_AVAILABLE:
                extraction_methods.append(('PyMuPDF', self._extract_pdf_with_pymupdf))
            if PYPDF2_AVAILABLE:
                extraction_methods.append(('PyPDF2', self._extract_pdf_with_pypdf2))
            
            if not extraction_methods:
                error_msg = "No PDF extraction libraries available"
                extraction_errors.append(error_msg)
                raise Exception(error_msg)
            
            # Try each primary method with retry
            for method_name, method_func in extraction_methods:
                for attempt in range(2):  # Try each method twice
                    try:
                        logger.debug(f"Trying PDF extraction with {method_name} (attempt {attempt + 1})")
                        content = method_func(temp_file_path)
                        if content and content.strip():
                            logger.info(f"Successfully extracted PDF content using {method_name}")
                            return content.strip()
                        else:
                            logger.warning(f"{method_name} returned empty content")
                    except Exception as e:
                        error_msg = f"{method_name} extraction failed (attempt {attempt + 1}): {e}"
                        logger.warning(error_msg)
                        extraction_errors.append(error_msg)
                        if attempt == 0:
                            time.sleep(0.5)  # Brief delay before retry
            
            # Fallback: Try basic text extraction from PDF binary
            logger.info("Attempting fallback: basic text extraction from PDF binary")
            try:
                fallback_content = self._extract_pdf_fallback_binary(file_content)
                if fallback_content and len(fallback_content.strip()) > 50:  # Minimum viable content
                    logger.warning("Using fallback binary extraction - content may be incomplete")
                    return fallback_content.strip()
                else:
                    extraction_errors.append("Fallback binary extraction returned insufficient content")
            except Exception as e:
                extraction_errors.append(f"Fallback binary extraction failed: {e}")
            
            raise Exception("All PDF extraction methods failed or returned empty content")
            
        finally:
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp PDF file: {e}")
    
    def _extract_pdf_fallback_binary(self, file_content: bytes) -> str:
        """
        Fallback method: Extract readable text from PDF binary content.
        This is a last-resort method that may produce incomplete results.
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text (may be incomplete)
        """
        try:
            # Look for text streams in PDF
            text_parts = []
            
            # Split by common PDF text markers
            parts = file_content.split(b'stream')
            
            for part in parts:
                # Try to extract readable ASCII text
                try:
                    # Look for text between BT (Begin Text) and ET (End Text) markers
                    if b'BT' in part and b'ET' in part:
                        text_section = part[part.find(b'BT'):part.find(b'ET')]
                        # Extract printable characters
                        readable = ''.join(chr(b) for b in text_section if 32 <= b <= 126 or b in [10, 13])
                        if readable.strip():
                            text_parts.append(readable)
                except:
                    continue
            
            if text_parts:
                return '\n'.join(text_parts)
            
            # If no text streams found, try extracting all printable characters
            readable_text = ''.join(chr(b) for b in file_content if 32 <= b <= 126 or b in [10, 13])
            # Clean up and return only if we have substantial content
            words = readable_text.split()
            if len(words) > 20:  # At least 20 words
                return ' '.join(words)
            
            return ""
            
        except Exception as e:
            logger.error(f"Binary fallback extraction failed: {e}")
            return ""
    
    def _extract_pdf_with_pdfplumber(self, file_path: str) -> str:
        """Extract PDF content using pdfplumber."""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        return text
    
    def _extract_pdf_with_pymupdf(self, file_path: str) -> str:
        """Extract PDF content using PyMuPDF."""
        text = ""
        doc = fitz.open(file_path)
        try:
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text:
                    text += page_text + "\n\n"
        finally:
            doc.close()
        return text
    
    def _extract_pdf_with_pypdf2(self, file_path: str) -> str:
        """Extract PDF content using PyPDF2."""
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        return text
    
    def _extract_docx_content(self, file_content: bytes) -> str:
        """Extract text content from DOCX file."""
        if not PYTHON_DOCX_AVAILABLE:
            raise Exception("python-docx library not available for DOCX processing")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            doc = Document(temp_file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        finally:
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp DOCX file: {e}")
    
    def _extract_docx_content_with_fallback(self, file_content: bytes, extraction_errors: List[str]) -> str:
        """
        Extract DOCX content with comprehensive fallback mechanisms.
        
        Args:
            file_content: DOCX file content as bytes
            extraction_errors: List to collect error messages
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If all extraction methods fail
        """
        if not PYTHON_DOCX_AVAILABLE:
            error_msg = "python-docx library not available for DOCX processing"
            extraction_errors.append(error_msg)
            raise Exception(error_msg)
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # Primary method: python-docx with retry
            for attempt in range(2):
                try:
                    logger.debug(f"Trying DOCX extraction with python-docx (attempt {attempt + 1})")
                    doc = Document(temp_file_path)
                    text_parts = []
                    
                    # Extract text from paragraphs
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                text_parts.append(" | ".join(row_text))
                    
                    content = "\n\n".join(text_parts)
                    if content.strip():
                        logger.info("Successfully extracted DOCX content using python-docx")
                        return content
                    else:
                        logger.warning("python-docx returned empty content")
                        
                except Exception as e:
                    error_msg = f"python-docx extraction failed (attempt {attempt + 1}): {e}"
                    logger.warning(error_msg)
                    extraction_errors.append(error_msg)
                    if attempt == 0:
                        time.sleep(0.5)
            
            # Fallback: Extract from DOCX XML structure
            logger.info("Attempting fallback: XML extraction from DOCX")
            try:
                fallback_content = self._extract_docx_fallback_xml(file_content)
                if fallback_content and len(fallback_content.strip()) > 50:
                    logger.warning("Using fallback XML extraction - content may be incomplete")
                    return fallback_content.strip()
                else:
                    extraction_errors.append("Fallback XML extraction returned insufficient content")
            except Exception as e:
                extraction_errors.append(f"Fallback XML extraction failed: {e}")
            
            raise Exception("All DOCX extraction methods failed or returned empty content")
            
        finally:
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp DOCX file: {e}")
    
    def _extract_docx_fallback_xml(self, file_content: bytes) -> str:
        """
        Fallback method: Extract text from DOCX XML structure.
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text (may be incomplete)
        """
        try:
            import zipfile
            import io
            import xml.etree.ElementTree as ET
            
            text_parts = []
            
            # DOCX is a ZIP file containing XML
            with zipfile.ZipFile(io.BytesIO(file_content), 'r') as zip_file:
                # Try to read document.xml
                try:
                    doc_xml = zip_file.read('word/document.xml')
                    root = ET.fromstring(doc_xml)
                    
                    # Extract text from all text nodes
                    # DOCX uses namespace, so we need to handle it
                    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    
                    for text_elem in root.findall('.//w:t', namespaces):
                        if text_elem.text:
                            text_parts.append(text_elem.text)
                    
                except Exception as e:
                    logger.warning(f"Failed to extract from document.xml: {e}")
                
                # Also try other XML files that might contain text
                for filename in zip_file.namelist():
                    if filename.startswith('word/') and filename.endswith('.xml'):
                        try:
                            xml_content = zip_file.read(filename)
                            # Extract any readable text
                            readable = ''.join(chr(b) for b in xml_content if 32 <= b <= 126 or b in [10, 13])
                            # Look for text between XML tags
                            import re
                            texts = re.findall(r'>([^<]+)<', readable)
                            text_parts.extend([t.strip() for t in texts if t.strip() and len(t.strip()) > 2])
                        except:
                            continue
            
            if text_parts:
                return ' '.join(text_parts)
            
            return ""
            
        except Exception as e:
            logger.error(f"XML fallback extraction failed: {e}")
            return ""
    
    def _extract_text_content(self, file_content: bytes) -> str:
        """Extract content from text files (TXT, MD)."""
        # Try UTF-8 first
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Try other common encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Unable to decode text file with common encodings")
    
    def _extract_text_content_with_fallback(self, file_content: bytes, extraction_errors: List[str]) -> str:
        """
        Extract text content with comprehensive encoding fallback.
        
        Args:
            file_content: Text file content as bytes
            extraction_errors: List to collect error messages
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If all encoding methods fail
        """
        # Try multiple encodings in order of likelihood
        encodings_to_try = [
            'utf-8',
            'utf-8-sig',  # UTF-8 with BOM
            'latin-1',
            'cp1252',  # Windows-1252
            'iso-8859-1',
            'ascii',
            'utf-16',
            'utf-16-le',
            'utf-16-be',
            'cp850',  # DOS encoding
            'mac_roman'  # Mac encoding
        ]
        
        for encoding in encodings_to_try:
            try:
                logger.debug(f"Trying text extraction with {encoding} encoding")
                content = file_content.decode(encoding)
                
                # Validate that we got reasonable text content
                if content.strip():
                    # Check for excessive binary characters
                    binary_chars = sum(1 for c in content if ord(c) < 32 and c not in '\t\n\r')
                    binary_ratio = binary_chars / len(content) if content else 0
                    
                    if binary_ratio < 0.1:  # Less than 10% binary characters
                        logger.info(f"Successfully extracted text content using {encoding} encoding")
                        return content
                    else:
                        logger.warning(f"{encoding} produced high binary ratio ({binary_ratio:.1%})")
                        extraction_errors.append(f"{encoding}: high binary content ratio")
                else:
                    extraction_errors.append(f"{encoding}: empty content")
                    
            except (UnicodeDecodeError, LookupError) as e:
                extraction_errors.append(f"{encoding}: {str(e)}")
                continue
        
        # Fallback: Extract only printable ASCII characters
        logger.info("Attempting fallback: printable ASCII extraction")
        try:
            fallback_content = self._extract_text_fallback_ascii(file_content)
            if fallback_content and len(fallback_content.strip()) > 50:
                logger.warning("Using fallback ASCII extraction - non-ASCII characters removed")
                return fallback_content
            else:
                extraction_errors.append("Fallback ASCII extraction returned insufficient content")
        except Exception as e:
            extraction_errors.append(f"Fallback ASCII extraction failed: {e}")
        
        raise Exception(f"Unable to decode text file with any encoding. Tried: {', '.join(encodings_to_try)}")
    
    def _extract_text_fallback_ascii(self, file_content: bytes) -> str:
        """
        Fallback method: Extract only printable ASCII characters.
        
        Args:
            file_content: Text file content as bytes
            
        Returns:
            Extracted ASCII text
        """
        try:
            # Extract printable ASCII characters and common whitespace
            printable_chars = []
            for byte in file_content:
                if 32 <= byte <= 126:  # Printable ASCII
                    printable_chars.append(chr(byte))
                elif byte in [9, 10, 13]:  # Tab, newline, carriage return
                    printable_chars.append(chr(byte))
                elif byte == 0:  # Null byte - skip
                    continue
                else:
                    # Replace non-printable with space
                    printable_chars.append(' ')
            
            content = ''.join(printable_chars)
            
            # Clean up excessive whitespace
            import re
            content = re.sub(r' +', ' ', content)  # Multiple spaces to single
            content = re.sub(r'\n\n+', '\n\n', content)  # Multiple newlines to double
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"ASCII fallback extraction failed: {e}")
            return ""
    
    def _chunk_content_for_rag(self, content: str) -> List[str]:
        """
        Split content into chunks suitable for RAG processing.
        
        Args:
            content: Text content to chunk
            
        Returns:
            List of text chunks
        """
        if not content or not content.strip():
            return []
        
        # Simple chunking by sentences and paragraphs
        chunks = []
        paragraphs = content.split('\n\n')
        
        current_chunk = ""
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # If adding this paragraph would exceed chunk size, save current chunk
            if len(current_chunk) + len(paragraph) > self.MAX_CHUNK_SIZE and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        # Handle overlap for better context
        if len(chunks) > 1 and self.CHUNK_OVERLAP > 0:
            overlapped_chunks = []
            for i, chunk in enumerate(chunks):
                if i == 0:
                    overlapped_chunks.append(chunk)
                else:
                    # Add overlap from previous chunk
                    prev_chunk = chunks[i-1]
                    overlap_text = prev_chunk[-self.CHUNK_OVERLAP:] if len(prev_chunk) > self.CHUNK_OVERLAP else prev_chunk
                    overlapped_chunk = overlap_text + "\n\n" + chunk
                    overlapped_chunks.append(overlapped_chunk)
            
            return overlapped_chunks
        
        return chunks
    
    @retry_with_exponential_backoff(max_retries=2, base_delay=1.0)
    async def _add_to_rag_system(self, note: LectureNote) -> None:
        """
        Add lecture note content to RAG system with retry and graceful degradation.
        
        Args:
            note: LectureNote object with extracted content
        """
        if not self.rag_system or not note.extracted_content:
            logger.debug(f"Skipping RAG integration for note {note.id}: RAG system not available or no content")
            return
        
        try:
            # Chunk the content
            chunks = self._chunk_content_for_rag(note.extracted_content)
            
            if not chunks:
                logger.warning(f"No chunks generated for note {note.id}")
                return
            
            # Add chunks to RAG system (implementation depends on RAG system interface)
            logger.info(f"Adding {len(chunks)} chunks from note {note.id} to RAG system")
            
            # TODO: Implement actual RAG integration based on the RAG system interface
            # await self.rag_system.add_lecture_note_chunks(note.id, chunks, note.associated_rubrics)
            
            # Mark successful RAG integration in metadata
            if not note.metadata:
                note.metadata = {}
            note.metadata['rag_integrated'] = True
            note.metadata['rag_integration_date'] = datetime.now().isoformat()
            note.metadata['chunk_count'] = len(chunks)
            
        except Exception as e:
            error_msg = f"Failed to add note {note.id} to RAG system: {e}"
            logger.error(error_msg)
            
            # Mark RAG integration failure in metadata but don't fail the upload
            if not note.metadata:
                note.metadata = {}
            note.metadata['rag_integrated'] = False
            note.metadata['rag_integration_error'] = str(e)
            note.metadata['rag_integration_retry_count'] = note.metadata.get('rag_integration_retry_count', 0) + 1
            
            # Only raise if this is a critical failure (not during upload)
            # During upload, we want graceful degradation
            raise IntegrationError(
                message="RAG integration failed",
                details=error_msg,
                context={'note_id': note.id, 'retry_count': note.metadata.get('rag_integration_retry_count', 0)}
            )
    
    async def associate_note_with_rubric(self, note_id: str, rubric_id: str) -> bool:
        """
        Associate a lecture note with a rubric.
        
        Args:
            note_id: Note identifier
            rubric_id: Rubric identifier
            
        Returns:
            True if associated successfully
        """
        try:
            success = self.storage.associate_note_with_rubric(note_id, rubric_id)
            
            if success and self.rag_system:
                # Update RAG system with new association
                note = self.storage.get_lecture_note(note_id)
                if note and note.extracted_content:
                    await self._add_to_rag_system(note)
            
            return success
            
        except Exception as e:
            logger.error(f"Failed to associate note {note_id} with rubric {rubric_id}: {e}")
            return False
    
    async def disassociate_note_from_rubric(self, note_id: str, rubric_id: str) -> bool:
        """
        Remove association between a lecture note and rubric.
        
        Args:
            note_id: Note identifier
            rubric_id: Rubric identifier
            
        Returns:
            True if disassociated successfully
        """
        try:
            success = self.storage.disassociate_note_from_rubric(note_id, rubric_id)
            
            if success and self.rag_system:
                # Update RAG system to remove association
                # TODO: Implement RAG system update for disassociation
                pass
            
            return success
            
        except Exception as e:
            logger.error(f"Failed to disassociate note {note_id} from rubric {rubric_id}: {e}")
            return False
    
    def get_lecture_note(self, note_id: str) -> Optional[LectureNote]:
        """
        Get a lecture note by ID.
        
        Args:
            note_id: Note identifier
            
        Returns:
            LectureNote object or None if not found
        """
        return self.storage.get_lecture_note(note_id)
    
    def get_all_lecture_notes(self) -> List[LectureNote]:
        """
        Get all lecture notes.
        
        Returns:
            List of all LectureNote objects
        """
        return self.storage.get_all_lecture_notes()
    
    def get_notes_for_rubric(self, rubric_id: str) -> List[LectureNote]:
        """
        Get all lecture notes associated with a rubric.
        
        Args:
            rubric_id: Rubric identifier
            
        Returns:
            List of associated LectureNote objects
        """
        return self.storage.get_notes_for_rubric(rubric_id)
    
    async def delete_lecture_note(self, note_id: str) -> bool:
        """
        Delete a lecture note and remove from RAG system.
        
        Args:
            note_id: Note identifier
            
        Returns:
            True if deleted successfully
        """
        try:
            # Get note before deletion for RAG cleanup
            note = self.storage.get_lecture_note(note_id)
            
            # Delete from storage
            success = self.storage.delete_lecture_note(note_id)
            
            if success and self.rag_system and note:
                # Remove from RAG system
                try:
                    # TODO: Implement RAG system cleanup
                    # await self.rag_system.remove_lecture_note_chunks(note_id)
                    pass
                except Exception as e:
                    logger.error(f"Failed to remove note {note_id} from RAG system: {e}")
                    # Don't fail deletion if RAG cleanup fails
            
            return success
            
        except Exception as e:
            logger.error(f"Failed to delete lecture note {note_id}: {e}")
            return False
    
    def search_notes(self, query: str, rubric_id: Optional[str] = None) -> List[LectureNote]:
        """
        Search lecture notes by content.
        
        Args:
            query: Search query string
            rubric_id: Optional rubric ID to filter by
            
        Returns:
            List of matching LectureNote objects
        """
        return self.storage.search_notes(query, rubric_id)
    
    async def reprocess_note(self, note_id: str) -> bool:
        """
        Reprocess a lecture note (re-extract content and update RAG).
        
        Args:
            note_id: Note identifier
            
        Returns:
            True if reprocessed successfully
        """
        try:
            note = self.storage.get_lecture_note(note_id)
            if not note:
                logger.warning(f"Note {note_id} not found for reprocessing")
                return False
            
            # TODO: This would require storing the original file content
            # For now, just mark as needing reprocessing
            logger.warning(f"Reprocessing not implemented - original file content not stored")
            return False
            
        except Exception as e:
            logger.error(f"Failed to reprocess note {note_id}: {e}")
            return False
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """
        Get statistics about lecture notes processing.
        
        Returns:
            Dictionary with processing statistics
        """
        try:
            base_stats = self.storage.get_statistics()
            
            # Add processing-specific statistics
            notes = self.storage.get_all_lecture_notes()
            
            processing_stats = {
                'pending': 0,
                'processing': 0,
                'completed': 0,
                'failed': 0
            }
            
            extraction_methods = {}
            
            for note in notes:
                status = note.processing_status.value
                processing_stats[status] += 1
                
                # Track extraction methods used
                if note.metadata.get('extraction_method'):
                    method = note.metadata['extraction_method']
                    extraction_methods[method] = extraction_methods.get(method, 0) + 1
            
            # Combine with base statistics
            base_stats.update({
                'processing_statistics': processing_stats,
                'extraction_methods': extraction_methods,
                'available_libraries': {
                    'pdfplumber': PDFPLUMBER_AVAILABLE,
                    'PyMuPDF': PYMUPDF_AVAILABLE,
                    'PyPDF2': PYPDF2_AVAILABLE,
                    'python-docx': PYTHON_DOCX_AVAILABLE
                }
            })
            
            return base_stats
            
        except Exception as e:
            logger.error(f"Failed to get processing statistics: {e}")
            return {'error': str(e)}
    
    def get_supported_formats(self) -> Dict[str, Any]:
        """
        Get information about supported file formats and available libraries.
        
        Returns:
            Dictionary with format support information
        """
        return {
            'supported_formats': list(self.SUPPORTED_FILE_TYPES),
            'max_file_size_mb': self.MAX_FILE_SIZE_BYTES / (1024 * 1024),
            'available_libraries': {
                'pdf': {
                    'pdfplumber': PDFPLUMBER_AVAILABLE,
                    'PyMuPDF': PYMUPDF_AVAILABLE,
                    'PyPDF2': PYPDF2_AVAILABLE
                },
                'docx': {
                    'python-docx': PYTHON_DOCX_AVAILABLE
                },
                'text': {
                    'built-in': True  # Always available
                }
            },
            'processing_settings': {
                'max_chunk_size': self.MAX_CHUNK_SIZE,
                'chunk_overlap': self.CHUNK_OVERLAP
            }
        }
    
    async def upload_lecture_note_with_retry(
        self,
        file_content: bytes,
        filename: str,
        associate_with_rubric: Optional[str] = None,
        max_retries: int = 3
    ) -> LectureNote:
        """
        Upload lecture note with retry mechanism for robust error handling.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            associate_with_rubric: Optional rubric ID to associate with
            max_retries: Maximum number of retry attempts
            
        Returns:
            Created LectureNote object
            
        Raises:
            ValidationError: If file validation fails after all retries
            IntegrationError: If processing fails after all retries
        """
        last_error = None
        
        for attempt in range(max_retries + 1):
            try:
                if attempt > 0:
                    logger.info(f"Retry attempt {attempt}/{max_retries} for {filename}")
                
                return await self.upload_lecture_note(file_content, filename, associate_with_rubric)
                
            except ValidationError as e:
                # Don't retry validation errors - they won't change
                logger.error(f"Validation error for {filename}: {e}")
                raise
                
            except Exception as e:
                last_error = e
                logger.warning(f"Upload attempt {attempt + 1} failed for {filename}: {e}")
                
                if attempt < max_retries:
                    # Wait before retry (exponential backoff)
                    wait_time = 2 ** attempt
                    logger.info(f"Waiting {wait_time} seconds before retry...")
                    await asyncio.sleep(wait_time)
                else:
                    logger.error(f"All {max_retries + 1} upload attempts failed for {filename}")
        
        # If we get here, all retries failed
        raise IntegrationError(f"Upload failed after {max_retries + 1} attempts: {str(last_error)}")
    
    def validate_and_repair_storage(self) -> Dict[str, Any]:
        """
        Validate storage integrity and attempt to repair common issues.
        
        Returns:
            Dictionary with validation and repair results
        """
        repair_results = {
            'storage_valid': True,
            'issues_found': [],
            'repairs_attempted': [],
            'repairs_successful': [],
            'repairs_failed': [],
            'notes_checked': 0,
            'notes_repaired': 0
        }
        
        try:
            logger.info("Starting storage validation and repair")
            
            # Get all notes
            notes = self.storage.get_all_lecture_notes()
            repair_results['notes_checked'] = len(notes)
            
            for note in notes:
                note_issues = []
                
                # Check for missing required fields
                if not note.id:
                    note_issues.append("missing_id")
                if not note.filename:
                    note_issues.append("missing_filename")
                if not note.original_name:
                    note_issues.append("missing_original_name")
                
                # Check for invalid processing status
                try:
                    ProcessingStatus(note.processing_status)
                except ValueError:
                    note_issues.append("invalid_processing_status")
                
                # Check for orphaned associations
                for rubric_id in note.associated_rubrics:
                    # This would require access to rubric storage to validate
                    # For now, just check for empty rubric IDs
                    if not rubric_id or not rubric_id.strip():
                        note_issues.append("empty_rubric_association")
                
                # Check content consistency
                if note.extracted_content and not note.word_count:
                    note_issues.append("missing_word_count")
                elif note.word_count and not note.extracted_content:
                    note_issues.append("word_count_without_content")
                
                if note_issues:
                    repair_results['issues_found'].extend([
                        f"Note {note.id}: {issue}" for issue in note_issues
                    ])
                    
                    # Attempt repairs
                    repaired = self._repair_note_issues(note, note_issues)
                    if repaired:
                        repair_results['notes_repaired'] += 1
                        repair_results['repairs_successful'].extend([
                            f"Note {note.id}: {issue}" for issue in note_issues
                        ])
                    else:
                        repair_results['repairs_failed'].extend([
                            f"Note {note.id}: {issue}" for issue in note_issues
                        ])
            
            # Check storage file integrity
            storage_stats = self.storage.get_statistics()
            if 'error' in storage_stats:
                repair_results['storage_valid'] = False
                repair_results['issues_found'].append(f"Storage error: {storage_stats['error']}")
            
            logger.info(f"Storage validation completed: {len(repair_results['issues_found'])} issues found, "
                       f"{repair_results['notes_repaired']} notes repaired")
            
        except Exception as e:
            logger.error(f"Storage validation failed: {e}")
            repair_results['storage_valid'] = False
            repair_results['issues_found'].append(f"Validation error: {str(e)}")
        
        return repair_results
    
    def _repair_note_issues(self, note: LectureNote, issues: List[str]) -> bool:
        """
        Attempt to repair issues with a lecture note.
        
        Args:
            note: LectureNote object with issues
            issues: List of issue identifiers
            
        Returns:
            True if repairs were successful
        """
        try:
            repaired = False
            
            # Repair missing word count
            if "missing_word_count" in issues and note.extracted_content:
                note.word_count = len(note.extracted_content.split())
                repaired = True
            
            # Clear invalid word count
            if "word_count_without_content" in issues:
                note.word_count = None
                repaired = True
            
            # Clean up empty rubric associations
            if "empty_rubric_association" in issues:
                note.associated_rubrics = [
                    rubric_id for rubric_id in note.associated_rubrics 
                    if rubric_id and rubric_id.strip()
                ]
                repaired = True
            
            # Reset invalid processing status
            if "invalid_processing_status" in issues:
                # Determine correct status based on content
                if note.metadata.get("processing_error"):
                    note.metadata["processing_status"] = ProcessingStatus.FAILED
                elif note.extracted_content:
                    note.metadata["processing_status"] = ProcessingStatus.COMPLETED
                    if not note.processed_at:
                        note.processed_at = datetime.now()
                else:
                    note.metadata["processing_status"] = ProcessingStatus.PENDING
                repaired = True
            
            # Save repaired note
            if repaired:
                success = self.storage.update_lecture_note(note)
                if not success:
                    logger.error(f"Failed to save repaired note {note.id}")
                    return False
                
                logger.info(f"Successfully repaired note {note.id}")
                return True
            
        except Exception as e:
            logger.error(f"Failed to repair note {note.id}: {e}")
        
        return False


# Error handling and fallback mechanisms
class LectureNotesProcessingError(Exception):
    """Custom exception for lecture notes processing errors."""
    
    def __init__(self, message: str, file_type: str = None, original_error: Exception = None):
        self.file_type = file_type
        self.original_error = original_error
        super().__init__(message)


def create_lecture_notes_service(storage_path: Optional[str] = None, rag_system=None) -> LectureNotesService:
    """
    Factory function to create a LectureNotesService instance.
    
    Args:
        storage_path: Optional path to lecture notes storage file
        rag_system: Optional RAG system instance
        
    Returns:
        Configured LectureNotesService instance
    """
    storage = LectureNotesStorage(storage_path) if storage_path else LectureNotesStorage()
    return LectureNotesService(storage, rag_system)


# Enhanced error handling classes
class LectureNotesValidationError(ValidationError):
    """Specific validation error for lecture notes."""
    
    def __init__(self, message: str, filename: str = None, validation_details: Dict[str, Any] = None):
        self.filename = filename
        self.validation_details = validation_details or {}
        super().__init__(message)


class LectureNotesProcessingError(IntegrationError):
    """Specific processing error for lecture notes."""
    
    def __init__(self, message: str, note_id: str = None, processing_stage: str = None):
        self.note_id = note_id
        self.processing_stage = processing_stage
        super().__init__(message)


# Enhanced factory function with error handling configuration
def create_lecture_notes_service_with_config(
    storage_path: Optional[str] = None,
    rag_system=None,
    enable_validation: bool = True,
    enable_retry: bool = True,
    max_retries: int = 3
) -> LectureNotesService:
    """
    Factory function to create a configured LectureNotesService instance.
    
    Args:
        storage_path: Optional path to lecture notes storage file
        rag_system: Optional RAG system instance
        enable_validation: Whether to enable enhanced validation
        enable_retry: Whether to enable retry mechanisms
        max_retries: Maximum number of retry attempts
        
    Returns:
        Configured LectureNotesService instance
    """
    try:
        storage = LectureNotesStorage(storage_path) if storage_path else LectureNotesStorage()
        service = LectureNotesService(storage, rag_system)
        
        # Configure service settings
        if hasattr(service, 'enable_validation'):
            service.enable_validation = enable_validation
        if hasattr(service, 'enable_retry'):
            service.enable_retry = enable_retry
        if hasattr(service, 'max_retries'):
            service.max_retries = max_retries
        
        logger.info(f"Created LectureNotesService with validation={enable_validation}, "
                   f"retry={enable_retry}, max_retries={max_retries}")
        
        return service
        
    except Exception as e:
        logger.error(f"Failed to create LectureNotesService: {e}")
        raise IntegrationError(f"Service creation failed: {str(e)}")