"""
Enhanced file validation and error handling for lecture notes.
Provides comprehensive validation for academic documents and robust error recovery.
"""

import os
import logging
import tempfile
import mimetypes
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
from datetime import datetime

# File processing imports with fallbacks
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import fitz  # pymupdf
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import magic
    PYTHON_MAGIC_AVAILABLE = True
except ImportError:
    PYTHON_MAGIC_AVAILABLE = False

try:
    from ..errors.grading_errors import ValidationError, ErrorHandler
except ImportError:
    from src.errors.grading_errors import ValidationError, ErrorHandler

logger = logging.getLogger(__name__)


class AcademicDocumentValidator:
    """Enhanced validator for academic documents with comprehensive error handling."""
    
    # File size limits
    MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50MB
    MIN_FILE_SIZE_BYTES = 10  # 10 bytes minimum
    
    # Supported file types with MIME types
    SUPPORTED_FORMATS = {
        'pdf': {
            'extensions': ['.pdf'],
            'mime_types': ['application/pdf'],
            'magic_bytes': [b'%PDF-'],
            'description': 'Portable Document Format'
        },
        'docx': {
            'extensions': ['.docx'],
            'mime_types': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
            'magic_bytes': [b'PK\x03\x04', b'PK\x05\x06', b'PK\x07\x08'],
            'description': 'Microsoft Word Document'
        },
        'txt': {
            'extensions': ['.txt'],
            'mime_types': ['text/plain'],
            'magic_bytes': [],  # Text files don't have specific magic bytes
            'description': 'Plain Text File'
        },
        'md': {
            'extensions': ['.md', '.markdown'],
            'mime_types': ['text/markdown', 'text/x-markdown'],
            'magic_bytes': [],  # Markdown files are text-based
            'description': 'Markdown Document'
        }
    }
    
    # Academic content indicators
    ACADEMIC_KEYWORDS = [
        'lecture', 'chapter', 'section', 'introduction', 'conclusion',
        'abstract', 'summary', 'references', 'bibliography', 'citation',
        'theory', 'concept', 'definition', 'example', 'exercise',
        'assignment', 'homework', 'quiz', 'exam', 'test', 'study',
        'research', 'analysis', 'methodology', 'results', 'discussion'
    ]
    
    def __init__(self):
        """Initialize the validator."""
        self._log_available_libraries()
    
    def _log_available_libraries(self) -> None:
        """Log available validation libraries."""
        available = []
        if PDFPLUMBER_AVAILABLE:
            available.append("pdfplumber")
        if PYMUPDF_AVAILABLE:
            available.append("PyMuPDF")
        if PYPDF2_AVAILABLE:
            available.append("PyPDF2")
        if PYTHON_DOCX_AVAILABLE:
            available.append("python-docx")
        if PYTHON_MAGIC_AVAILABLE:
            available.append("python-magic")
        
        logger.info(f"Available validation libraries: {', '.join(available) if available else 'Basic only'}")
    
    def validate_file_comprehensive(
        self,
        file_content: bytes,
        filename: str,
        expected_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Perform comprehensive file validation with enhanced error handling.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            expected_type: Expected file type (optional)
            
        Returns:
            Dictionary with detailed validation results
        """
        validation_result = {
            'valid': False,
            'filename': filename,
            'file_size': len(file_content),
            'detected_type': None,
            'format_validation': {},
            'content_validation': {},
            'academic_indicators': {},
            'errors': [],
            'warnings': [],
            'fallback_options': []
        }
        
        try:
            # Basic size validation
            size_validation = self._validate_file_size(file_content)
            if not size_validation['valid']:
                validation_result['errors'].extend(size_validation['errors'])
                return validation_result
            
            # File type detection and validation
            type_validation = self._detect_and_validate_file_type(file_content, filename, expected_type)
            validation_result.update(type_validation)
            
            if not type_validation['valid']:
                return validation_result
            
            detected_type = type_validation['detected_type']
            
            # Format-specific validation
            format_validation = self._validate_file_format_enhanced(file_content, detected_type)
            validation_result['format_validation'] = format_validation
            
            if not format_validation.get('valid', False):
                validation_result['errors'].extend(format_validation.get('errors', []))
                # Try fallback methods
                fallback_result = self._try_fallback_validation(file_content, detected_type)
                validation_result['fallback_options'] = fallback_result
                
                if not any(fb['valid'] for fb in fallback_result):
                    return validation_result
            
            # Content validation for academic documents
            content_validation = self._validate_academic_content(file_content, detected_type)
            validation_result['content_validation'] = content_validation
            
            # Academic indicators analysis
            if content_validation.get('extracted_content'):
                academic_analysis = self._analyze_academic_indicators(content_validation['extracted_content'])
                validation_result['academic_indicators'] = academic_analysis
            
            # Final validation decision
            validation_result['valid'] = (
                format_validation.get('valid', False) or
                any(fb['valid'] for fb in validation_result['fallback_options'])
            )
            
            return validation_result
            
        except Exception as e:
            logger.error(f"Comprehensive validation failed: {e}")
            validation_result['errors'].append(f"Validation error: {str(e)}")
            return validation_result
    
    def _validate_file_size(self, file_content: bytes) -> Dict[str, Any]:
        """Validate file size constraints."""
        file_size = len(file_content)
        
        if file_size < self.MIN_FILE_SIZE_BYTES:
            return {
                'valid': False,
                'errors': [f"File too small ({file_size} bytes). Minimum size is {self.MIN_FILE_SIZE_BYTES} bytes."]
            }
        
        if file_size > self.MAX_FILE_SIZE_BYTES:
            return {
                'valid': False,
                'errors': [f"File too large ({file_size / (1024*1024):.1f}MB). Maximum size is {self.MAX_FILE_SIZE_BYTES / (1024*1024):.0f}MB."]
            }
        
        return {'valid': True, 'file_size': file_size}
    
    def _detect_and_validate_file_type(
        self,
        file_content: bytes,
        filename: str,
        expected_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """Enhanced file type detection using multiple methods."""
        result = {
            'valid': False,
            'detected_type': None,
            'detection_methods': {},
            'errors': [],
            'warnings': []
        }
        
        # Method 1: File extension
        file_ext = Path(filename).suffix.lower()
        ext_detected_type = None
        for file_type, info in self.SUPPORTED_FORMATS.items():
            if file_ext in info['extensions']:
                ext_detected_type = file_type
                break
        
        result['detection_methods']['extension'] = {
            'detected_type': ext_detected_type,
            'confidence': 'high' if ext_detected_type else 'none'
        }
        
        # Method 2: Magic bytes
        magic_detected_type = self._detect_by_magic_bytes(file_content)
        result['detection_methods']['magic_bytes'] = {
            'detected_type': magic_detected_type,
            'confidence': 'high' if magic_detected_type else 'none'
        }
        
        # Method 3: MIME type (if python-magic available)
        mime_detected_type = None
        if PYTHON_MAGIC_AVAILABLE:
            try:
                mime_type = magic.from_buffer(file_content, mime=True)
                for file_type, info in self.SUPPORTED_FORMATS.items():
                    if mime_type in info['mime_types']:
                        mime_detected_type = file_type
                        break
            except Exception as e:
                logger.warning(f"MIME type detection failed: {e}")
        
        result['detection_methods']['mime_type'] = {
            'detected_type': mime_detected_type,
            'confidence': 'medium' if mime_detected_type else 'none'
        }
        
        # Determine final detected type
        detected_types = [ext_detected_type, magic_detected_type, mime_detected_type]
        detected_types = [dt for dt in detected_types if dt is not None]
        
        if not detected_types:
            result['errors'].append("Unable to detect file type using any method")
            return result
        
        # Check for consistency
        if len(set(detected_types)) == 1:
            result['detected_type'] = detected_types[0]
            result['valid'] = True
        else:
            # Conflicting detections - use priority order
            for priority_type in [magic_detected_type, ext_detected_type, mime_detected_type]:
                if priority_type:
                    result['detected_type'] = priority_type
                    result['valid'] = True
                    result['warnings'].append(f"Conflicting type detections: {detected_types}. Using {priority_type}")
                    break
        
        # Validate against expected type if provided
        if expected_type and result['detected_type'] != expected_type:
            result['warnings'].append(f"Detected type '{result['detected_type']}' differs from expected '{expected_type}'")
        
        # Check if detected type is supported
        if result['detected_type'] not in self.SUPPORTED_FORMATS:
            result['valid'] = False
            result['errors'].append(f"File type '{result['detected_type']}' is not supported")
        
        return result
    
    def _detect_by_magic_bytes(self, file_content: bytes) -> Optional[str]:
        """Detect file type by magic bytes."""
        for file_type, info in self.SUPPORTED_FORMATS.items():
            for magic_bytes in info['magic_bytes']:
                if file_content.startswith(magic_bytes):
                    return file_type
        return None
    
    def _validate_file_format_enhanced(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        """Enhanced format-specific validation with multiple fallback methods."""
        if file_type == 'pdf':
            return self._validate_pdf_enhanced(file_content)
        elif file_type == 'docx':
            return self._validate_docx_enhanced(file_content)
        elif file_type in ['txt', 'md']:
            return self._validate_text_enhanced(file_content)
        else:
            return {'valid': False, 'errors': [f'Unsupported file type: {file_type}']}
    
    def _validate_pdf_enhanced(self, file_content: bytes) -> Dict[str, Any]:
        """Enhanced PDF validation with multiple library fallbacks."""
        result = {
            'valid': False,
            'errors': [],
            'warnings': [],
            'metadata': {},
            'library_results': {}
        }
        
        # Create temporary file for validation
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # Try each available PDF library
            libraries = []
            if PDFPLUMBER_AVAILABLE:
                libraries.append(('pdfplumber', self._validate_pdf_with_pdfplumber))
            if PYMUPDF_AVAILABLE:
                libraries.append(('PyMuPDF', self._validate_pdf_with_pymupdf))
            if PYPDF2_AVAILABLE:
                libraries.append(('PyPDF2', self._validate_pdf_with_pypdf2))
            
            if not libraries:
                result['errors'].append("No PDF processing libraries available")
                return result
            
            successful_validations = []
            
            for lib_name, lib_func in libraries:
                try:
                    lib_result = lib_func(temp_file_path)
                    result['library_results'][lib_name] = lib_result
                    
                    if lib_result.get('valid', False):
                        successful_validations.append(lib_name)
                        # Merge metadata from successful validation
                        if 'metadata' in lib_result:
                            result['metadata'].update(lib_result['metadata'])
                
                except Exception as e:
                    result['library_results'][lib_name] = {
                        'valid': False,
                        'error': str(e)
                    }
                    logger.warning(f"PDF validation with {lib_name} failed: {e}")
            
            if successful_validations:
                result['valid'] = True
                result['successful_libraries'] = successful_validations
            else:
                result['errors'].append("All PDF validation methods failed")
                
        finally:
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp PDF file: {e}")
        
        return result
    
    def _validate_pdf_with_pdfplumber(self, file_path: str) -> Dict[str, Any]:
        """Validate PDF using pdfplumber."""
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                # Check if we can extract text from at least one page
                text_extractable = False
                for i, page in enumerate(pdf.pages[:3]):  # Check first 3 pages
                    if page.extract_text():
                        text_extractable = True
                        break
                
                metadata = {
                    'page_count': page_count,
                    'text_extractable': text_extractable,
                    'library': 'pdfplumber'
                }
                
                # Get PDF metadata if available
                if hasattr(pdf, 'metadata') and pdf.metadata:
                    metadata['pdf_metadata'] = dict(pdf.metadata)
                
                return {
                    'valid': True,
                    'metadata': metadata
                }
                
        except Exception as e:
            return {
                'valid': False,
                'error': f"pdfplumber validation failed: {str(e)}"
            }
    
    def _validate_pdf_with_pymupdf(self, file_path: str) -> Dict[str, Any]:
        """Validate PDF using PyMuPDF."""
        try:
            doc = fitz.open(file_path)
            
            try:
                page_count = doc.page_count
                
                # Check if we can extract text
                text_extractable = False
                for page_num in range(min(3, page_count)):  # Check first 3 pages
                    page = doc[page_num]
                    if page.get_text().strip():
                        text_extractable = True
                        break
                
                metadata = {
                    'page_count': page_count,
                    'text_extractable': text_extractable,
                    'library': 'PyMuPDF'
                }
                
                # Get PDF metadata
                if doc.metadata:
                    metadata['pdf_metadata'] = doc.metadata
                
                return {
                    'valid': True,
                    'metadata': metadata
                }
                
            finally:
                doc.close()
                
        except Exception as e:
            return {
                'valid': False,
                'error': f"PyMuPDF validation failed: {str(e)}"
            }
    
    def _validate_pdf_with_pypdf2(self, file_path: str) -> Dict[str, Any]:
        """Validate PDF using PyPDF2."""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                page_count = len(reader.pages)
                
                # Check if we can extract text
                text_extractable = False
                for i in range(min(3, page_count)):  # Check first 3 pages
                    page = reader.pages[i]
                    if page.extract_text().strip():
                        text_extractable = True
                        break
                
                metadata = {
                    'page_count': page_count,
                    'text_extractable': text_extractable,
                    'library': 'PyPDF2'
                }
                
                # Get PDF metadata
                if reader.metadata:
                    metadata['pdf_metadata'] = dict(reader.metadata)
                
                return {
                    'valid': True,
                    'metadata': metadata
                }
                
        except Exception as e:
            return {
                'valid': False,
                'error': f"PyPDF2 validation failed: {str(e)}"
            }
    
    def _validate_docx_enhanced(self, file_content: bytes) -> Dict[str, Any]:
        """Enhanced DOCX validation."""
        result = {
            'valid': False,
            'errors': [],
            'warnings': [],
            'metadata': {}
        }
        
        if not PYTHON_DOCX_AVAILABLE:
            result['errors'].append("python-docx library not available for DOCX validation")
            return result
        
        # Check ZIP signature
        if not file_content.startswith(b'PK'):
            result['errors'].append("Invalid DOCX format: missing ZIP signature")
            return result
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            doc = Document(temp_file_path)
            
            # Count content elements
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            # Check for actual content
            has_text_content = any(p.text.strip() for p in doc.paragraphs)
            
            metadata = {
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'has_text_content': has_text_content,
                'library': 'python-docx'
            }
            
            # Get document properties if available
            try:
                core_props = doc.core_properties
                if core_props:
                    metadata['document_properties'] = {
                        'title': core_props.title,
                        'author': core_props.author,
                        'subject': core_props.subject,
                        'created': core_props.created.isoformat() if core_props.created else None,
                        'modified': core_props.modified.isoformat() if core_props.modified else None
                    }
            except Exception as e:
                logger.debug(f"Could not extract document properties: {e}")
            
            result['valid'] = True
            result['metadata'] = metadata
            
            if not has_text_content:
                result['warnings'].append("Document appears to have no readable text content")
            
        except Exception as e:
            result['errors'].append(f"DOCX validation failed: {str(e)}")
        
        finally:
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp DOCX file: {e}")
        
        return result
    
    def _validate_text_enhanced(self, file_content: bytes) -> Dict[str, Any]:
        """Enhanced text file validation with encoding detection."""
        result = {
            'valid': False,
            'errors': [],
            'warnings': [],
            'metadata': {}
        }
        
        # Try multiple encodings
        encodings_to_try = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings_to_try:
            try:
                text_content = file_content.decode(encoding)
                
                # Analyze content
                line_count = len(text_content.splitlines())
                word_count = len(text_content.split())
                char_count = len(text_content)
                
                # Check for binary content indicators
                binary_indicators = sum(1 for c in text_content if ord(c) < 32 and c not in '\t\n\r')
                binary_ratio = binary_indicators / char_count if char_count > 0 else 0
                
                metadata = {
                    'encoding': encoding,
                    'line_count': line_count,
                    'word_count': word_count,
                    'character_count': char_count,
                    'binary_ratio': binary_ratio,
                    'has_content': bool(text_content.strip())
                }
                
                # Validate content quality
                if binary_ratio > 0.1:  # More than 10% binary characters
                    result['warnings'].append(f"High binary content ratio ({binary_ratio:.1%}) - may not be a text file")
                
                if not text_content.strip():
                    result['warnings'].append("File appears to be empty or contain only whitespace")
                
                result['valid'] = True
                result['metadata'] = metadata
                break
                
            except UnicodeDecodeError:
                continue
        
        if not result['valid']:
            result['errors'].append(f"Unable to decode text file with any of these encodings: {', '.join(encodings_to_try)}")
        
        return result
    
    def _try_fallback_validation(self, file_content: bytes, file_type: str) -> List[Dict[str, Any]]:
        """Try fallback validation methods when primary validation fails."""
        fallback_results = []
        
        if file_type == 'pdf':
            # Try alternative PDF validation approaches
            fallback_results.extend(self._pdf_fallback_methods(file_content))
        elif file_type == 'docx':
            # Try alternative DOCX validation approaches
            fallback_results.extend(self._docx_fallback_methods(file_content))
        elif file_type in ['txt', 'md']:
            # Try alternative text validation approaches
            fallback_results.extend(self._text_fallback_methods(file_content))
        
        return fallback_results
    
    def _pdf_fallback_methods(self, file_content: bytes) -> List[Dict[str, Any]]:
        """Fallback methods for PDF validation."""
        fallbacks = []
        
        # Fallback 1: Basic PDF structure check
        try:
            if b'%PDF-' in file_content and b'%%EOF' in file_content:
                fallbacks.append({
                    'method': 'basic_structure_check',
                    'valid': True,
                    'confidence': 'low',
                    'description': 'PDF structure markers found'
                })
            else:
                fallbacks.append({
                    'method': 'basic_structure_check',
                    'valid': False,
                    'description': 'PDF structure markers not found'
                })
        except Exception as e:
            fallbacks.append({
                'method': 'basic_structure_check',
                'valid': False,
                'error': str(e)
            })
        
        # Fallback 2: Try as binary data with manual parsing
        try:
            # Look for common PDF objects
            pdf_objects = [b'/Type', b'/Page', b'/Font', b'/Contents']
            found_objects = sum(1 for obj in pdf_objects if obj in file_content)
            
            if found_objects >= 2:
                fallbacks.append({
                    'method': 'object_detection',
                    'valid': True,
                    'confidence': 'medium',
                    'description': f'Found {found_objects} PDF object markers'
                })
            else:
                fallbacks.append({
                    'method': 'object_detection',
                    'valid': False,
                    'description': f'Only found {found_objects} PDF object markers'
                })
        except Exception as e:
            fallbacks.append({
                'method': 'object_detection',
                'valid': False,
                'error': str(e)
            })
        
        return fallbacks
    
    def _docx_fallback_methods(self, file_content: bytes) -> List[Dict[str, Any]]:
        """Fallback methods for DOCX validation."""
        fallbacks = []
        
        # Fallback 1: ZIP structure check
        try:
            import zipfile
            import io
            
            with zipfile.ZipFile(io.BytesIO(file_content), 'r') as zip_file:
                file_list = zip_file.namelist()
                
                # Check for essential DOCX files
                essential_files = ['word/document.xml', '[Content_Types].xml']
                found_essential = sum(1 for ef in essential_files if ef in file_list)
                
                if found_essential >= 1:
                    fallbacks.append({
                        'method': 'zip_structure_check',
                        'valid': True,
                        'confidence': 'medium',
                        'description': f'Found {found_essential} essential DOCX files'
                    })
                else:
                    fallbacks.append({
                        'method': 'zip_structure_check',
                        'valid': False,
                        'description': 'Essential DOCX files not found in ZIP structure'
                    })
                    
        except Exception as e:
            fallbacks.append({
                'method': 'zip_structure_check',
                'valid': False,
                'error': str(e)
            })
        
        return fallbacks
    
    def _text_fallback_methods(self, file_content: bytes) -> List[Dict[str, Any]]:
        """Fallback methods for text validation."""
        fallbacks = []
        
        # Fallback 1: Treat as binary and look for text patterns
        try:
            # Count printable ASCII characters
            printable_count = sum(1 for b in file_content if 32 <= b <= 126 or b in [9, 10, 13])
            printable_ratio = printable_count / len(file_content) if file_content else 0
            
            if printable_ratio > 0.7:  # More than 70% printable characters
                fallbacks.append({
                    'method': 'printable_ratio_check',
                    'valid': True,
                    'confidence': 'medium',
                    'description': f'High printable character ratio ({printable_ratio:.1%})'
                })
            else:
                fallbacks.append({
                    'method': 'printable_ratio_check',
                    'valid': False,
                    'description': f'Low printable character ratio ({printable_ratio:.1%})'
                })
                
        except Exception as e:
            fallbacks.append({
                'method': 'printable_ratio_check',
                'valid': False,
                'error': str(e)
            })
        
        return fallbacks
    
    def _validate_academic_content(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        """Validate content for academic document characteristics."""
        result = {
            'extracted_content': None,
            'content_length': 0,
            'extraction_method': None,
            'errors': [],
            'warnings': []
        }
        
        try:
            # Extract a sample of content for analysis
            if file_type == 'pdf':
                content = self._extract_pdf_sample(file_content)
            elif file_type == 'docx':
                content = self._extract_docx_sample(file_content)
            elif file_type in ['txt', 'md']:
                content = self._extract_text_sample(file_content)
            else:
                result['errors'].append(f"Content extraction not supported for {file_type}")
                return result
            
            if content:
                result['extracted_content'] = content
                result['content_length'] = len(content)
                result['extraction_method'] = f"{file_type}_sample"
            else:
                result['warnings'].append("No content could be extracted for analysis")
            
        except Exception as e:
            result['errors'].append(f"Content extraction failed: {str(e)}")
        
        return result
    
    def _extract_pdf_sample(self, file_content: bytes) -> Optional[str]:
        """Extract a sample of text from PDF for analysis."""
        # Try the most reliable method available
        if PDFPLUMBER_AVAILABLE:
            return self._extract_pdf_sample_pdfplumber(file_content)
        elif PYMUPDF_AVAILABLE:
            return self._extract_pdf_sample_pymupdf(file_content)
        elif PYPDF2_AVAILABLE:
            return self._extract_pdf_sample_pypdf2(file_content)
        return None
    
    def _extract_pdf_sample_pdfplumber(self, file_content: bytes) -> Optional[str]:
        """Extract PDF sample using pdfplumber."""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            with pdfplumber.open(temp_file_path) as pdf:
                # Extract text from first few pages
                text_parts = []
                for i, page in enumerate(pdf.pages[:3]):  # First 3 pages
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                return "\n\n".join(text_parts)[:2000]  # Limit to 2000 chars
        except Exception:
            return None
        finally:
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass
    
    def _extract_pdf_sample_pymupdf(self, file_content: bytes) -> Optional[str]:
        """Extract PDF sample using PyMuPDF."""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            doc = fitz.open(temp_file_path)
            text_parts = []
            
            for page_num in range(min(3, doc.page_count)):  # First 3 pages
                page = doc[page_num]
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)
            
            doc.close()
            return "\n\n".join(text_parts)[:2000]  # Limit to 2000 chars
        except Exception:
            return None
        finally:
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass
    
    def _extract_pdf_sample_pypdf2(self, file_content: bytes) -> Optional[str]:
        """Extract PDF sample using PyPDF2."""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            with open(temp_file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text_parts = []
                
                for i in range(min(3, len(reader.pages))):  # First 3 pages
                    page = reader.pages[i]
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                return "\n\n".join(text_parts)[:2000]  # Limit to 2000 chars
        except Exception:
            return None
        finally:
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass
    
    def _extract_docx_sample(self, file_content: bytes) -> Optional[str]:
        """Extract a sample of text from DOCX for analysis."""
        if not PYTHON_DOCX_AVAILABLE:
            return None
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            doc = Document(temp_file_path)
            text_parts = []
            
            # Extract text from first few paragraphs
            for i, paragraph in enumerate(doc.paragraphs[:10]):  # First 10 paragraphs
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            return "\n\n".join(text_parts)[:2000]  # Limit to 2000 chars
        except Exception:
            return None
        finally:
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass
    
    def _extract_text_sample(self, file_content: bytes) -> Optional[str]:
        """Extract a sample of text from text files."""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text_content = file_content.decode(encoding)
                return text_content[:2000]  # Limit to 2000 chars
            except UnicodeDecodeError:
                continue
        
        return None
    
    def _analyze_academic_indicators(self, content: str) -> Dict[str, Any]:
        """Analyze content for academic document indicators."""
        if not content:
            return {'score': 0, 'indicators': [], 'confidence': 'none'}
        
        content_lower = content.lower()
        
        # Count academic keywords
        found_keywords = []
        for keyword in self.ACADEMIC_KEYWORDS:
            if keyword in content_lower:
                found_keywords.append(keyword)
        
        # Calculate academic score
        keyword_score = len(found_keywords) / len(self.ACADEMIC_KEYWORDS)
        
        # Check for academic structure indicators
        structure_indicators = []
        
        # Look for common academic document structures
        if any(pattern in content_lower for pattern in ['chapter', 'section', 'subsection']):
            structure_indicators.append('hierarchical_structure')
        
        if any(pattern in content_lower for pattern in ['references', 'bibliography', 'citations']):
            structure_indicators.append('references')
        
        if any(pattern in content_lower for pattern in ['abstract', 'introduction', 'conclusion']):
            structure_indicators.append('academic_sections')
        
        if any(pattern in content_lower for pattern in ['figure', 'table', 'equation']):
            structure_indicators.append('academic_elements')
        
        # Calculate overall academic confidence
        structure_score = len(structure_indicators) / 4  # 4 possible structure types
        overall_score = (keyword_score * 0.6) + (structure_score * 0.4)
        
        if overall_score > 0.3:
            confidence = 'high'
        elif overall_score > 0.1:
            confidence = 'medium'
        else:
            confidence = 'low'
        
        return {
            'score': overall_score,
            'keyword_score': keyword_score,
            'structure_score': structure_score,
            'found_keywords': found_keywords,
            'structure_indicators': structure_indicators,
            'confidence': confidence,
            'content_sample_length': len(content)
        }
    
    def get_validation_summary(self, validation_result: Dict[str, Any]) -> str:
        """Generate a human-readable validation summary."""
        if not validation_result:
            return "No validation results available"
        
        summary_parts = []
        
        # Basic file info
        filename = validation_result.get('filename', 'Unknown')
        file_size = validation_result.get('file_size', 0)
        detected_type = validation_result.get('detected_type', 'Unknown')
        
        summary_parts.append(f"File: {filename}")
        summary_parts.append(f"Size: {file_size / 1024:.1f} KB")
        summary_parts.append(f"Type: {detected_type}")
        
        # Validation status
        if validation_result.get('valid', False):
            summary_parts.append("✓ Validation: PASSED")
        else:
            summary_parts.append("✗ Validation: FAILED")
        
        # Errors and warnings
        errors = validation_result.get('errors', [])
        warnings = validation_result.get('warnings', [])
        
        if errors:
            summary_parts.append(f"Errors: {len(errors)}")
            for error in errors[:3]:  # Show first 3 errors
                summary_parts.append(f"  - {error}")
        
        if warnings:
            summary_parts.append(f"Warnings: {len(warnings)}")
            for warning in warnings[:2]:  # Show first 2 warnings
                summary_parts.append(f"  - {warning}")
        
        # Academic indicators
        academic_indicators = validation_result.get('academic_indicators', {})
        if academic_indicators:
            confidence = academic_indicators.get('confidence', 'unknown')
            score = academic_indicators.get('score', 0)
            summary_parts.append(f"Academic content: {confidence} confidence ({score:.1%})")
        
        return "\n".join(summary_parts)


# Factory function
def create_academic_validator() -> AcademicDocumentValidator:
    """Create an instance of the academic document validator."""
    return AcademicDocumentValidator()